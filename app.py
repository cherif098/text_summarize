# app.py
import streamlit as st
import PyPDF2
import docx
import spacy
from deep_translator import GoogleTranslator
import pandas as pd
from docx import Document
import io
from spacy.cli import download
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from heapq import nlargest
from string import punctuation
import numpy as np

# Téléchargement des modèles spaCy pour différentes langues
for lang in ['fr_core_news_sm', 'en_core_web_sm', 'de_core_news_sm', 'es_core_news_sm']:
    try:
        spacy.load(lang)
    except:
        download(lang)

@st.cache_resource
def load_spacy_models():
    return {
        'fr': spacy.load('fr_core_news_sm'),
        'en': spacy.load('en_core_web_sm'),
        'de': spacy.load('de_core_news_sm'),
        'es': spacy.load('es_core_news_sm')
    }

def detect_language(text):
    first_words = text.lower().split()[:10]
    fr_words = set(['le', 'la', 'les', 'un', 'une', 'des', 'et', 'est'])
    en_words = set(['the', 'a', 'an', 'and', 'is', 'are', 'were'])
    de_words = set(['der', 'die', 'das', 'und', 'ist'])
    es_words = set(['el', 'la', 'los', 'las', 'un', 'una', 'y', 'es'])
    
    lang_scores = {
        'fr': len(set(first_words) & fr_words),
        'en': len(set(first_words) & en_words),
        'de': len(set(first_words) & de_words),
        'es': len(set(first_words) & es_words)
    }
    return max(lang_scores, key=lang_scores.get)

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def generate_summary(text, precision, target_lang, nlp_models):
    # Détecter la langue source
    source_lang = detect_language(text)
    
    # Traduire en anglais si nécessaire
    if source_lang != 'en':
        translator = GoogleTranslator(source=source_lang, target='en')
        text = translator.translate(text)
    
    # Utiliser le modèle spaCy anglais pour l'analyse
    doc = nlp_models['en'](text)
    
    # Calculer les scores de fréquence des mots
    word_freq = {}
    for word in doc:
        if word.text.lower() not in nlp_models['en'].Defaults.stop_words and word.text.lower() not in punctuation:
            if word.text not in word_freq.keys():
                word_freq[word.text] = 1
            else:
                word_freq[word.text] += 1

    # Normaliser les fréquences
    max_freq = max(word_freq.values())
    for word in word_freq.keys():
        word_freq[word] = word_freq[word] / max_freq

    # Calculer les scores des phrases
    sent_scores = {}
    for sent in doc.sents:
        for word in sent:
            if word.text in word_freq.keys():
                if sent not in sent_scores.keys():
                    sent_scores[sent] = word_freq[word.text]
                else:
                    sent_scores[sent] += word_freq[word.text]

    # Sélectionner les meilleures phrases selon la précision
    select_length = {
        'précis': 0.3,    # 30% du texte original
        'moyen': 0.2,     # 20% du texte original
        'vague': 0.1      # 10% du texte original
    }
    
    select_k = max(1, int(len(list(doc.sents)) * select_length[precision]))
    summary_sentences = nlargest(select_k, sent_scores, key=sent_scores.get)
    summary = ' '.join([str(s) for s in summary_sentences])
    
    # Traduire vers la langue cible si nécessaire
    if target_lang != 'en':
        translator = GoogleTranslator(source='en', target=target_lang)
        summary = translator.translate(summary)
    
    return summary

def analyze_text(text, lang, nlp_models):
    doc = nlp_models[lang](text)
    analysis = []
    
    for sent in doc.sents:
        subjects = []
        verbs = []
        complements = []
        
        for token in sent:
            if "subj" in token.dep_:
                subjects.append(token.text)
            elif token.pos_ == "VERB":
                verbs.append(token.text)
            elif "obj" in token.dep_ or token.dep_ == "iobj":
                complements.append(token.text)
                
        analysis.append({
            'Sujet': ' '.join(subjects) if subjects else '',
            'Verbe': ' '.join(verbs) if verbs else '',
            'Complément': ' '.join(complements) if complements else ''
        })
    
    return pd.DataFrame(analysis)

def create_pdf(text, font_family, font_size, style):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    font_name = {
        'Arial': 'Helvetica',
        'Times New Roman': 'Times-Roman',
        'Calibri': 'Helvetica'
    }.get(font_family, 'Helvetica')
    
    if style == 'bold':
        font_name += '-Bold'
    elif style == 'italic':
        font_name += '-Oblique'
    
    text_object = c.beginText()
    text_object.setFont(font_name, font_size)
    text_object.setTextOrigin(50, height - 50)
    
    words = text.split()
    lines = []
    current_line = []
    max_width = width - 100
    
    for word in words:
        current_line.append(word)
        line_width = c.stringWidth(' '.join(current_line), font_name, font_size)
        if line_width > max_width:
            current_line.pop()
            lines.append(' '.join(current_line))
            current_line = [word]
    
    if current_line:
        lines.append(' '.join(current_line))
    
    for line in lines:
        text_object.textLine(line)
    
    c.drawText(text_object)
    c.save()
    buffer.seek(0)
    return buffer

def save_as_docx(text, font_family, font_size, style):
    doc = Document()
    paragraph = doc.add_paragraph(text)
    
    run = paragraph.runs[0]
    run.font.name = font_family
    run.font.size = docx.shared.Pt(int(font_size))
    if style == 'italic':
        run.italic = True
    elif style == 'bold':
        run.bold = True
    
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def main():
    st.title("Générateur de Résumés Multilingue")
    st.markdown("""
    <style>
    .stApp {
        background-color: #f5f5f5;
    }
    .stButton button {
        background-color: #4CAF50;
        color: white;
        padding: 10px 20px;
        border-radius: 5px;
    }
    .stSelectbox select {
        border-radius: 5px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Charger les modèles spaCy
    nlp_models = load_spacy_models()
    
    input_method = st.radio("Choisissez votre méthode d'entrée:", 
                           ["Texte", "Fichier PDF", "Fichier Word"])
    
    text = None
    if input_method == "Texte":
        text = st.text_area("Entrez votre texte ici:", height=200)
    elif input_method == "Fichier PDF":
        uploaded_file = st.file_uploader("Choisissez un fichier PDF", type="pdf")
        if uploaded_file:
            text = extract_text_from_pdf(uploaded_file)
            st.success("Fichier PDF chargé avec succès!")
    else:
        uploaded_file = st.file_uploader("Choisissez un fichier Word", type="docx")
        if uploaded_file:
            text = extract_text_from_docx(uploaded_file)
            st.success("Fichier Word chargé avec succès!")
    
    col1, col2 = st.columns(2)
    with col1:
        precision = st.selectbox("Niveau de précision:", 
                               ["précis", "moyen", "vague"])
    with col2:
        target_lang = st.selectbox("Langue cible:", 
                                 ["français", "anglais", "allemand", "espagnol"])
    
    st.subheader("Options de style pour l'exportation")
    col3, col4, col5 = st.columns(3)
    with col3:
        font_family = st.selectbox("Police:", 
                                 ["Arial", "Times New Roman", "Calibri"])
    with col4:
        font_size = st.selectbox("Taille:", 
                               [10, 11, 12, 14, 16, 18])
    with col5:
        style = st.selectbox("Style:", 
                           ["normal", "bold", "italic"])
    
    if st.button("Générer le résumé"):
        if text and text.strip():
            lang_map = {
                "français": "fr",
                "anglais": "en",
                "allemand": "de",
                "espagnol": "es"
            }
            
            with st.spinner("Génération du résumé en cours..."):
                try:
                    summary = generate_summary(text, precision, lang_map[target_lang], nlp_models)
                    st.subheader("Résumé généré:")
                    st.write(summary)
                    
                    st.subheader("Analyse grammaticale:")
                    analysis_df = analyze_text(summary, lang_map[target_lang], nlp_models)
                    st.dataframe(analysis_df.style.set_properties(**{
                        'background-color': 'lightblue',
                        'color': 'black',
                        'border-color': 'white'
                    }))
                    
                    st.subheader("Télécharger le résumé")
                    col6, col7 = st.columns(2)
                    
                    with col6:
                        pdf_buffer = create_pdf(summary, font_family, font_size, style)
                        st.download_button(
                            "Télécharger en PDF",
                            pdf_buffer,
                            "resume.pdf",
                            "application/pdf"
                        )
                    
                    with col7:
                        docx_buffer = save_as_docx(summary, font_family, font_size, style)
                        st.download_button(
                            "Télécharger en Word",
                            docx_buffer,
                            "resume.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Une erreur est survenue: {str(e)}")
        else:
            st.error("Veuillez entrer du texte ou charger un fichier.")

if __name__ == "__main__":
    main()